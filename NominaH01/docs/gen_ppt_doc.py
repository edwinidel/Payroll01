from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

root = Path(__file__).resolve().parents[1]
output_path = root / 'docs' / 'presentacion-planilla.docx'
output_path.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

logo_candidates = [
    root / '2FA' / 'wwwroot' / 'img' / 'logo.png',
    root / '2FA' / 'wwwroot' / 'img' / 'logo1.png',
]
logo_path = next((p for p in logo_candidates if p.exists()), None)
if logo_path:
    doc.add_picture(str(logo_path), width=Inches(2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

heading = doc.add_heading('Solucion Integral de Planilla y Portal del Empleado', 0)
heading.alignment = 1

p = doc.add_paragraph()
p.add_run('Propuesta de valor').bold = True
doc.add_paragraph('Automatizamos planilla, ausencias y produccion destajo, integrando control de asistencia GPS y autoservicio para empleados.')

doc.add_heading('Alcance Modular', level=1)
modules = [
    'Core de planilla: catalogos, formulas, periodos, pre y post nomina.',
    'Control de asistencia: turnos, marcaciones, excepciones y aprobaciones.',
    'Produccion destajo: captura de unidades, factores, centros de costo y aprobaciones.',
    'Deducciones y obligaciones: legales, embargos, terceros.',
    'Reportes y exportaciones: banca, fondos, analitica.',
    'Portal del empleado: talonarios, certificados, balances, notificaciones push/email.',
    'App movil: GPS, selfie opcional, modo offline y sincronizacion.',
]
for item in modules:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('App Movil (Asistencia y Destajo)', level=1)
app_points = [
    'Flujo de fichaje con ubicacion GPS, validacion de radio y selfie opcional.',
    'Trabajo online/offline con cola de sincronizacion y marca de tiempo segura.',
    'Captura de destajo: actividad, unidad, cantidad, centro de costo y evidencia foto.',
    'Aprobaciones por supervisor y reglas por proyecto/ubicacion.',
    'Alertas: marca fuera de radio, doble marcacion, horas excedidas.',
]
for item in app_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Portal del Empleado', level=1)
portal_points = [
    'Talonarios historicos con descarga PDF.',
    'Solicitudes de certificado (laboral, ingresos) con generacion automatica.',
    'Consulta de balances: vacaciones, dias personales, horas extra.',
    'Bandeja de notificaciones y reglas por rol/ubicacion.',
    'Soporte web responsive y app movil unificada.',
]
for item in portal_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Integraciones', level=1)
integrations = [
    'Bancos locales: generacion de archivos de pago.',
    'ERP/contabilidad: export de asientos y centros de costo.',
    'SSO/LDAP opcional y MFA/2FA ya habilitado.',
    'APIs REST para asistencia, destajo y certificados.',
]
for item in integrations:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Plan de Entrega', level=1)
plan_points = [
    'Fase 1 (4 semanas): puesta en marcha planilla, catalogos y banca.',
    'Fase 2 (4 semanas): asistencia, turnos y portal del empleado (web).',
    'Fase 3 (4 semanas): app movil GPS + destajo y notificaciones.',
    'Acompanamiento: capacitacion, manuales y soporte local.',
]
for item in plan_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Beneficios Clave', level=1)
benefits = [
    'Menos reprocesos: reglas claras y validaciones automaticas.',
    'Transparencia al empleado: autoservicio y trazabilidad.',
    'Cumplimiento: deducciones legales y auditoria.',
    'Escalabilidad: arquitectura cloud-ready y APIs abiertas.',
]
for item in benefits:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Siguientes Pasos', level=1)
next_steps = [
    'Demo guiada con escenarios de asistencia GPS y destajo.',
    'Validar archivos de banco y plan contable.',
    'Definir reglas de proyectos, radios y jerarquias de aprobacion.',
]
for item in next_steps:
    doc.add_paragraph(item, style='List Bullet')

doc.save(output_path)
print(f'Generated {output_path}')
